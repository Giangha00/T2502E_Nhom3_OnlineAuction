#!/usr/bin/env python3
"""Generate Canva AI presentation DOCX from OnlineAuction project analysis."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "CardMarket_OnlineAuction_Canva_Presentation.docx"


def set_run_font(run, size=11, bold=False, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def add_heading_custom(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    sizes = {0: 22, 1: 16, 2: 13, 3: 12}
    set_run_font(run, size=sizes.get(level, 12), bold=True, color=RGBColor(0x1A, 0x1A, 0x2E))
    p.paragraph_format.space_before = Pt(14 if level <= 1 else 8)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_label(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=11, bold=True, color=RGBColor(0xB4, 0x53, 0x09))
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        set_run_font(run, size=11)
        p.paragraph_format.space_after = Pt(2)


def add_plain(doc, text, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=11)
    run.italic = italic
    p.paragraph_format.space_after = Pt(4)
    return p


def add_separator(doc):
    p = doc.add_paragraph()
    run = p.add_run("─" * 48)
    set_run_font(run, size=9, color=RGBColor(0x99, 0x99, 0x99))
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)


def add_slide_block(doc, number, title, bullets, image_suggestion, speaker_note, extras=None):
    add_heading_custom(doc, f"# Slide {number}", level=1)
    add_heading_custom(doc, title, level=2)
    if bullets:
        add_bullets(doc, bullets)
    if extras:
        for block in extras:
            kind = block.get("kind")
            if kind == "label":
                add_label(doc, block["text"])
            elif kind == "bullets":
                add_bullets(doc, block["items"])
            elif kind == "plain":
                add_plain(doc, block["text"], italic=block.get("italic", False))
            elif kind == "code":
                p = doc.add_paragraph()
                run = p.add_run(block["text"])
                set_run_font(run, size=10)
                run.font.name = "Consolas"
                p.paragraph_format.space_after = Pt(4)
            elif kind == "table":
                rows = block["rows"]
                table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                table.style = "Table Grid"
                for i, row in enumerate(rows):
                    for j, cell_text in enumerate(row):
                        cell = table.rows[i].cells[j]
                        cell.text = ""
                        p = cell.paragraphs[0]
                        run = p.add_run(cell_text)
                        set_run_font(run, size=10, bold=(i == 0))
                doc.add_paragraph()
    add_label(doc, "Image Suggestion:")
    add_plain(doc, image_suggestion)
    add_label(doc, "Speaker Note:")
    add_plain(doc, speaker_note)
    add_separator(doc)


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Pt(56)
    section.bottom_margin = Pt(56)
    section.left_margin = Pt(64)
    section.right_margin = Pt(64)

    # Cover
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("PRESENTATION DOCUMENT")
    set_run_font(run, size=20, bold=True, color=RGBColor(0x1A, 0x1A, 0x2E))

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("CardMarket — OnlineAuction (Nhóm 3)")
    set_run_font(run, size=14, bold=True)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(
        "Tài liệu phục vụ Canva AI tạo slide thuyết trình\n"
        "Nguồn: source code + docs trong repository (không bịa nội dung)"
    )
    set_run_font(run, size=10, color=RGBColor(0x55, 0x55, 0x55))

    add_plain(
        doc,
        "Hướng dẫn Canva AI: Mỗi mục # Slide = 1 slide. Giữ 3–7 bullet/slide. "
        "Chèn ảnh theo Image Placeholder / đường dẫn local. Speaker Note dùng cho ghi chú thuyết trình.",
        italic=True,
    )
    add_separator(doc)

    # Slide 1
    add_slide_block(
        doc,
        1,
        "Giới thiệu đề tài",
        [
            "Tên đề tài: Hệ thống đấu giá trực tuyến CardMarket (OnlineAuction)",
            "Mục tiêu: Marketplace đấu giá & mua ngay collectibles",
            "Đối tượng: Bidder, Seller, Admin vận hành",
            "Tính năng chính: Verify listing, PayPal, realtime bid, FCM",
            "Ý nghĩa: Giao dịch thẻ TCG/thể thao minh bạch, an toàn",
        ],
        "Hero homepage CardMarket + thẻ PSA / sports cards. "
        "Chèn ảnh: _screenshots/01_home.png",
        "Nêu tên dự án OnlineAuction/CardMarket (Nhóm 3). Nhấn 3 đối tượng User–Seller–Admin. Chuyển nhanh sang thành viên.",
    )

    # Slide 2
    add_slide_block(
        doc,
        2,
        "Giới thiệu thành viên",
        None,
        "Lưới 6 avatar team. Chèn ảnh: OnlineAuction/wwwroot/images/team/*.png",
        "Nguồn: Views/AboutUs/About.cshtml + SharedResource. Nếu cần role dev thật, bổ sung MSSV ngoài slide này.",
        extras=[
            {
                "kind": "table",
                "rows": [
                    ["Tên", "Vai trò (About)", "Công việc"],
                    ["Phạm Việt Anh", "Founder & CEO", "Tầm nhìn sản phẩm, chiến lược marketplace"],
                    ["Danil Fomin Long", "Head of curation", "Chất lượng listing, chuẩn xác thực"],
                    ["Nguyễn Văn Hưng", "Technical director", "Kiến trúc, bidding, payment, verification"],
                    ["Đinh Văn Hải", "Marketplace operations", "Onboarding seller, vận hành live auction"],
                    ["Nguyễn Hữu Quân", "Trust & compliance", "Chống gian lận, xác minh listing"],
                    ["Nguyễn Giang Hà", "Growth & community", "Cộng đồng collector, hỗ trợ seller/bidder"],
                ],
            }
        ],
    )

    # Slide 3
    add_slide_block(
        doc,
        3,
        "Mục lục",
        [
            "Giới thiệu bài toán",
            "Chức năng người dùng",
            "Chức năng Admin",
            "Activity Diagram",
            "Kiến trúc hệ thống",
            "Database & Công nghệ",
            "Kết luận & Cảm ơn",
        ],
        "Timeline / numbered agenda tối giản, nền tối, không icon rối.",
        "Dùng làm roadmap — chỉ đọc tiêu đề, không giải thích sâu.",
    )

    # Slide 4
    add_slide_block(
        doc,
        4,
        "Giới thiệu bài toán",
        [
            "Bối cảnh: Collector cần sàn đấu giá thẻ tin cậy",
            "Vấn đề: Listing giả, thiếu realtime bid, thanh toán rời rạc",
            "Giải pháp: ASP.NET Core MVC + Identity + PayPal + FCM",
            "Lợi ích: Verify listing, bid live, order center, admin permissions",
        ],
        "Split before/after: _screenshots/04_auction_marketplace.png + 05_auction_detail.png",
        "Một câu về grading PSA/BGS (About). Nhấn luồng: đăng bán → admin verify → bid → thắng → thanh toán.",
    )

    # Slide 5
    add_slide_block(
        doc,
        5,
        "Chức năng người dùng",
        None,
        "Collage auth + bid + order: _screenshots/05_login.png, 14_place_bid_view.png, 11_orders.png",
        "Nêu URL chính. COD đánh dấu paid ngay; PayPal qua capture return.",
        extras=[
            {"kind": "label", "text": "Authentication"},
            {
                "kind": "bullets",
                "items": [
                    "Đăng ký, xác nhận email, đăng nhập/đăng xuất",
                    "Quên mật khẩu OTP; dual cookie tách Admin",
                ],
            },
            {"kind": "label", "text": "Auction"},
            {
                "kind": "bullets",
                "items": [
                    "Duyệt /Auction, /BuyNow, chi tiết & lịch sử bid",
                    "Đăng ký phiên + đặt cọc PayPal (nếu bắt buộc)",
                    "Đặt bid (rate limit / fraud / anti-snipe)",
                    "Tạo listing /Sell → trạng thái confirming",
                ],
            },
            {"kind": "label", "text": "Profile"},
            {
                "kind": "bullets",
                "items": [
                    "Account: bids, watchlist, selling, orders",
                    "Sửa hồ sơ; xem seller profile",
                ],
            },
            {"kind": "label", "text": "Notification"},
            {
                "kind": "bullets",
                "items": [
                    "In-app dropdown + FCM web push",
                    "Outbid, sắp kết thúc, thắng, thanh toán",
                ],
            },
            {"kind": "label", "text": "Order · Payment · Shipping"},
            {
                "kind": "bullets",
                "items": [
                    "Payment Center /Order (auction_win + buy_now)",
                    "PayPal Sandbox hoặc COD",
                    "Form giao hàng: tên, địa chỉ, city, phone",
                ],
            },
        ],
    )

    # Slide 6
    add_slide_block(
        doc,
        6,
        "Chức năng Admin",
        None,
        "Admin UI: _screenshots/16_admin_dashboard.png → 19_admin_verify.png → 23_admin_complaints.png",
        "Nhấn dual session: /Admin/Account/Login cookie .AuctionHouse.Admin.",
        extras=[
            {"kind": "label", "text": "Dashboard"},
            {"kind": "bullets", "items": ["Metrics + filter ngày; export report"]},
            {"kind": "label", "text": "User Management"},
            {"kind": "bullets", "items": ["Xem/sửa user; fraud alerts liên quan"]},
            {"kind": "label", "text": "Auction Management"},
            {"kind": "bullets", "items": ["CRUD phiên; tạo live bypass review"]},
            {"kind": "label", "text": "Verify Auctions"},
            {"kind": "bullets", "items": ["Approve / Reject listing confirming"]},
            {"kind": "label", "text": "Product · Category · Buy Now"},
            {
                "kind": "bullets",
                "items": [
                    "Catalog products & templates; CRUD categories",
                    "Quản lý listing buy_now",
                ],
            },
            {"kind": "label", "text": "Complaints · Permissions"},
            {
                "kind": "bullets",
                "items": [
                    "Review refund/complaint",
                    "Permission động (RequirePermission); Admin bypass",
                ],
            },
            {
                "kind": "plain",
                "text": "Ghi chú: Không có Admin Order Management / Notification module riêng trong Controllers.",
                "italic": True,
            },
        ],
    )

    # Slide 7
    add_slide_block(
        doc,
        7,
        "Activity Diagram",
        [
            "Không vẽ diagram trong tài liệu này",
            "Chỉ placeholder + mô tả để Canva AI / chèn ảnh sẵn",
        ],
        "Horizontal swimlane activity. Ưu tiên chèn ảnh có sẵn trong repo.",
        "Walk-through một lần theo mũi tên. Không đi sâu anti-fraud trừ khi được hỏi.",
        extras=[
            {"kind": "label", "text": "Image Placeholder — Activity Diagram"},
            {
                "kind": "code",
                "text": "User Register → Login → Browse Product → Bid → Win Auction → Payment → Shipping → Complete",
            },
            {"kind": "label", "text": "Chi tiết luồng (từ code/docs)"},
            {
                "kind": "bullets",
                "items": [
                    "SignUp / Confirm email → Login (User scheme)",
                    "Browse /Auction → Detail → Register (+ deposit PayPal nếu cần)",
                    "Place bid → SignalR BidUpdated",
                    "AuctionFinalizationWorker tạo order auction_win",
                    "/Order nhập shipping → PayPal/COD → Confirmation",
                ],
            },
            {"kind": "label", "text": "Chèn ảnh (vị trí đặt)"},
            {
                "kind": "bullets",
                "items": [
                    "_diagram_check/ACT_4_2_1_Register_Account_Flow_*.png",
                    "_diagram_check/ACT_4_2_4_Place_Bid_Flow_*.png",
                    "_diagram_check/ACT_4_2_6_Payment_Flow_*.png",
                    "_diagram_gen/act_register_auction.png (meta.json key act_register_auction)",
                ],
            },
            {
                "kind": "plain",
                "text": "Use Case tham chiếu: docs/use-case-diagram.md + _diagram_check/UC_4_3_*.png",
                "italic": True,
            },
        ],
    )

    # Slide 8
    add_slide_block(
        doc,
        8,
        "Kiến trúc hệ thống",
        [
            "Client: Razor Views + Tailwind + SignalR JS",
            "ASP.NET Core MVC 8 Controllers / Admin Area",
            "Service Layer: Auction, Bid, Order, PayPal, Notification…",
            "SQL Server (default) / MySQL optional + Cloudinary",
            "RabbitMQ (email + auction lifecycle) + Firebase FCM",
            "PayPal API + Gmail API; Azure App Service (Production)",
            "Workers: AuctionFinalizationWorker, RabbitMqConsumer",
        ],
        "Layer diagram trái→phải theo mô tả Architecture Diagram bên dưới. Repo không có PNG Architecture riêng.",
        "Nhấn dual cookie Identity + permission policies Admin. Realtime = SignalR; offline push = FCM.",
        extras=[
            {"kind": "label", "text": "Image Placeholder — Architecture Diagram"},
            {
                "kind": "code",
                "text": (
                    "[Browser]\n"
                    "   ↓\n"
                    "[ASP.NET Core MVC 8]\n"
                    "   ↓\n"
                    "[Service Layer]\n"
                    "   ↓\n"
                    "[SQL Server / MySQL] + [Cloudinary]\n"
                    "   ↓\n"
                    "[RabbitMQ] [Firebase FCM] [PayPal] [Gmail]\n"
                    "[Azure App Service] [Background Workers]"
                ),
            },
            {
                "kind": "plain",
                "text": "Nguồn: Program.cs, OnlineAuction.csproj, OnlineAuction/README.md",
                "italic": True,
            },
        ],
    )

    # Slide 9
    add_slide_block(
        doc,
        9,
        "Database",
        [
            "Không vẽ ERD trong tài liệu — chỉ placeholder + bảng chính",
            "Default provider: SqlServer (appsettings.json DatabaseProvider)",
        ],
        "Chèn Class/ERD: _class_gen/class_diagram.png (PlantUML URL trong _class_gen/url.txt).",
        "Highlight Users → Products → Auctions → Bids → Orders → Payments. Fraud/deposit nhắc một câu.",
        extras=[
            {"kind": "label", "text": "Image Placeholder — Database Diagram / ERD"},
            {
                "kind": "plain",
                "text": "Database Link (dbdiagram.io): Không tìm thấy link dbdiagram.io trong repository.",
                "italic": True,
            },
            {
                "kind": "plain",
                "text": "Schema SQL tham chiếu: OnlineAuction/Data/Scripts/rarecard_schema.sql",
                "italic": True,
            },
            {"kind": "label", "text": "Bảng chính (AuctionHouseDbContext)"},
            {
                "kind": "table",
                "rows": [
                    ["Nhóm", "Tables"],
                    ["Identity", "users, roles, user_roles, claims/logins/tokens"],
                    ["Catalog", "categories, product_templates, products, product_images, product_documents"],
                    ["Auction", "auctions, bids, auction_registrations, auction_registration_deposits"],
                    ["Commerce", "orders, order_items, payments"],
                    ["Engage", "notifications, user_device_tokens, watchlist_items"],
                    ["Trust", "complaints, bid_fraud_alerts, winner_non_payment_logs"],
                    ["Other", "user_otp_codes, user_sandbox_wallets"],
                ],
            },
        ],
    )

    # Slide 10
    add_slide_block(
        doc,
        10,
        "Công nghệ sử dụng",
        None,
        "Logo strip tech stack trên nền tối (ASP.NET, SQL Server, RabbitMQ, Firebase, Cloudinary, PayPal, Azure).",
        "Không liệt kê version package trừ khi ban giám khảo hỏi.",
        extras=[
            {
                "kind": "table",
                "rows": [
                    ["Layer", "Công nghệ (trong repo)"],
                    ["Backend", "ASP.NET Core 8 MVC, EF Core 9, hosted workers"],
                    ["Frontend", "Razor Views, Tailwind CSS, jQuery, SignalR client"],
                    ["Database", "SQL Server (default); MySQL (Pomelo) / SQLite packages"],
                    ["Authentication", "ASP.NET Core Identity, dual cookie schemes"],
                    ["Storage", "Cloudinary (ảnh/avatar)"],
                    ["Queue", "RabbitMQ.Client (email + auction lifecycle)"],
                    ["Messaging", "Firebase Admin (FCM Web Push)"],
                    ["Payment", "PayPal REST Sandbox (+ COD)"],
                    ["Deployment", "Azure App Service settings (Production)"],
                    ["Tools", "Rider/dotnet CLI, xUnit, Playwright E2E, ClosedXML, Bogus"],
                ],
            }
        ],
    )

    # Slide 11
    add_slide_block(
        doc,
        11,
        "Kết luận",
        None,
        "Checkmark roadmap 4 cột: Achieved / Hard / Strength / Next.",
        "Kết bằng demo URL localhost:5006 hoặc Azure nếu đã deploy.",
        extras=[
            {"kind": "label", "text": "Mục tiêu đạt được"},
            {
                "kind": "bullets",
                "items": [
                    "Marketplace auction + buy now end-to-end",
                    "Admin verify, permissions, dashboard reports",
                    "PayPal checkout, notification, realtime bid",
                ],
            },
            {"kind": "label", "text": "Khó khăn"},
            {
                "kind": "bullets",
                "items": [
                    "Dual session Admin/User",
                    "Đồng bộ order/payment/deposit & non-payment recovery",
                    "Fraud/rate-limit + anti-snipe khi bid cao",
                ],
            },
            {"kind": "label", "text": "Ưu điểm"},
            {
                "kind": "bullets",
                "items": [
                    "Service layer rõ; permission động",
                    "SignalR + FCM; smoke/E2E tests",
                ],
            },
            {"kind": "label", "text": "Hướng phát triển"},
            {
                "kind": "bullets",
                "items": [
                    "Google OAuth (docs identity có sẵn)",
                    "Seller rating thật từ DB",
                    "PayPal live + monitoring capture recovery",
                ],
            },
        ],
    )

    # Slide 12
    add_slide_block(
        doc,
        12,
        "Cảm ơn",
        [
            "Cảm ơn Quý thầy cô và các bạn đã lắng nghe",
            "CardMarket — đấu giá collectibles minh bạch và tin cậy",
        ],
        "Full-bleed homepage hero mờ + chữ Cảm ơn lớn. Chèn: _screenshots/01_home.png",
        "Mời hỏi đáp. Demo: user1@auctionhouse.local / Admin admin@auctionhouse.com (README).",
    )

    # Appendix
    add_heading_custom(doc, "Phụ lục — Asset paths cho Canva", level=1)
    add_bullets(
        doc,
        [
            "Screenshots UI: _screenshots/*.png",
            "Activity diagrams: _diagram_gen/*.png, _diagram_check/ACT_*.png",
            "Use case diagrams: _diagram_check/UC_*.png, docs/use-case-diagram.md",
            "Class/ERD: _class_gen/class_diagram.png",
            "Sequence: _seq_gen/*.puml",
            "dbdiagram.io: không có trong source",
        ],
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    build()
